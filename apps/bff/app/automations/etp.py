# apps/bff/app/automations/etp.py
from __future__ import annotations

"""
Automação ETP — Estudo Técnico Preliminar.

Propósito
---------
- Receber um formulário simples de ETP, validar os campos principais,
  persistir a submissão e gerar um documento final a partir do modelo DOCX.
- Preservar cabeçalho e rodapé do arquivo-base `model.docx`.
- Expor endpoints de schema, submissão, consulta, download e UI HTML.

Estratégia de geração DOCX
--------------------------
- Quando o template possuir placeholders Jinja/docxtpl, utiliza `render_docx_template`.
- Quando o template estiver em branco (somente com cabeçalho/rodapé), monta o corpo
  do ETP programaticamente com `python-docx`, preservando o timbre institucional.

Segurança/RBAC
--------------
- Submeter e listar exige o papel "compras".
- Downloads aceitam também "admin".
- Auditoria administrativa exige "admin".

Persistência
------------
- Usa `submissions` e `automation_audits` via `app.db`.
"""

import json
import logging
import mimetypes
import os
import pathlib
import psycopg
import re
from datetime import datetime
from io import BytesIO
from typing import Any, Dict, List, Optional
from uuid import uuid4

from fastapi import APIRouter, BackgroundTasks, Depends, HTTPException, Request
from pydantic import BaseModel, ConfigDict, Field, ValidationError, field_validator, model_validator
from starlette.responses import HTMLResponse, StreamingResponse

from app.auth.rbac import require_roles_any
from app.db import add_audit, get_submission, insert_submission, list_audits, list_submissions, update_submission, DATABASE_URL
from app.notifications import send_notification
from app.utils.docx_tools import convert_docx_to_pdf, get_docx_placeholders, render_docx_template

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Pt
except Exception:  # pragma: no cover
    Document = None
    WD_ALIGN_PARAGRAPH = None
    qn = None
    Pt = None

logger = logging.getLogger(__name__)

KIND = "etp"
ETP_VERSION = "1.0.0"
TITLE = "ETP — Estudo Técnico Preliminar"
AUTOMATION_META = {
    "kind": KIND,
    "version": ETP_VERSION,
    "title": TITLE,
}
REQUIRED_ROLES = ("compras",)
ELEVATED_ROLES = ("admin",)

TEMPLATE_PATH = pathlib.Path(os.environ.get("ETP_TEMPLATE_PATH", "/app/templates/etp/model.docx"))
OUTPUT_DIR = pathlib.Path(os.environ.get("ETP_OUTPUT_DIR", "/app/data/files/etp"))
TPL_DIR = pathlib.Path(__file__).resolve().parent / "templates" / "etp"

MAX_SHORT = 300
MAX_MEDIUM = 2000
MAX_LONG = 12000
EMAIL_RE = re.compile(r"^[^@\s]+@[^@\s]+\.[^@\s]+$")
LEGACY_SOLUCAO_FIELD_RE = re.compile(r"^descricaosolucao(?:\d+)?$", re.IGNORECASE)
LEGACY_ANALISE_FIELD_RE = re.compile(r"^analisesolucao(?:\d+)?$", re.IGNORECASE)


router = APIRouter(prefix=f"/api/automations/{KIND}", tags=[f"automation:{KIND}"])


def err_json(status: int, **payload: Any) -> StreamingResponse:
    """
    Retorna JSON UTF-8 com mensagens amigáveis.
    """
    return StreamingResponse(
        BytesIO(json.dumps(payload, ensure_ascii=False).encode("utf-8")),
        status_code=status,
        media_type="application/json; charset=utf-8",
    )


def _to_obj(x: Any, default: Any = None) -> Any:
    """
    Converte dict/list/bytes/str JSON para objeto Python.
    """
    if x is None:
        return {} if default is None else default
    if isinstance(x, (dict, list)):
        return x
    if isinstance(x, (bytes, bytearray)):
        try:
            return json.loads(x.decode("utf-8"))
        except Exception:
            return {} if default is None else default
    if isinstance(x, str):
        try:
            return json.loads(x)
        except Exception:
            return {} if default is None else default
    return {} if default is None else default


def _safe_comp(txt: str, *, max_len: int = 80) -> str:
    """
    Normaliza componentes de filename removendo caracteres perigosos.
    """
    s = re.sub(r"[^A-Za-z0-9._-]+", "_", str(txt or "")).strip("._-")
    return (s or "arquivo")[:max_len]


def _normalize_text(v: Any) -> str:
    """
    Converte valores arbitrários em texto limpo.
    """
    if v is None:
        return ""
    if isinstance(v, str):
        return v.strip()
    if isinstance(v, (int, float, bool)):
        return str(v).strip()
    return str(v).strip()


def _normalize_string_list(v: Any) -> List[str]:
    """
    Aceita lista, texto único ou itens em formato dict e devolve lista limpa.

    Importante:
    - preserva quebras de linha dentro de cada item textual;
    - não transforma tópicos/Alt+Enter em novos itens da lista;
    - mantém compatibilidade com listas e dicionários legados.
    """
    out: List[str] = []

    def _push(x: Any) -> None:
        if x is None:
            return
        if isinstance(x, str):
            txt = _normalize_text(x)
            if txt:
                out.append(txt)
            return
        if isinstance(x, dict):
            for key in ("texto", "descricao", "nome", "titulo", "label", "value"):
                val = _normalize_text(x.get(key))
                if val:
                    out.append(val)
                    return
            flat = _normalize_text(x)
            if flat:
                out.append(flat)
            return
        if isinstance(x, list):
            for item in x:
                _push(item)
            return
        flat = _normalize_text(x)
        if flat:
            out.append(flat)

    _push(v)
    return out


def _normalize_quantitativos(v: Any) -> List[Dict[str, str]]:
    """
    Normaliza a tabela de quantitativos aceitando diferentes nomes de campo.
    """
    items = v if isinstance(v, list) else []
    out: List[Dict[str, str]] = []

    for idx, raw in enumerate(items, start=1):
        if isinstance(raw, str):
            txt = _normalize_text(raw)
            if txt:
                out.append(
                    {
                        "item": str(idx).zfill(2),
                        "especificacao": txt,
                        "unidade": "Unidade",
                        "quantidade": "1",
                    }
                )
            continue

        if not isinstance(raw, dict):
            continue

        item_no = _normalize_text(raw.get("item") or raw.get("numero") or raw.get("ordem") or str(idx).zfill(2))
        especificacao = _normalize_text(
            raw.get("especificacao") or raw.get("descricao") or raw.get("objeto") or raw.get("texto")
        )
        unidade = _normalize_text(raw.get("unidade") or raw.get("unidadeMedida") or raw.get("medida") or "Unidade")
        quantidade = _normalize_text(raw.get("quantidade") or raw.get("qtd") or raw.get("valor") or "1")

        has_any = bool(item_no or especificacao or unidade or quantidade)
        if not has_any:
            continue

        out.append(
            {
                "item": item_no or str(idx).zfill(2),
                "especificacao": especificacao,
                "unidade": unidade or "Unidade",
                "quantidade": quantidade or "1",
            }
        )

    return out


def _coalesce_string_list(body: Dict[str, Any], *keys: str) -> List[str]:
    """
    Retorna a primeira lista textual não vazia dentre as chaves informadas.
    """
    for key in keys:
        value = body.get(key)
        items = _normalize_string_list(value)
        if items:
            return items
    return []


def _has_any_role(user: Dict[str, Any], *roles: str) -> bool:
    """
    Verifica se o usuário possui ao menos um dos papéis informados.
    """
    user_roles = set((user or {}).get("roles") or [])
    return any(r in user_roles for r in roles)


def _owns_submission(row: Dict[str, Any], user: Dict[str, Any]) -> bool:
    """
    Verifica se a submissão pertence ao usuário autenticado.
    """
    u_cpf = (user.get("cpf") or "").strip() or None
    u_email = (user.get("email") or "").strip() or None
    owner_cpf = (row.get("actor_cpf") or "").strip() or None
    owner_email = (row.get("actor_email") or "").strip() or None
    return bool(
        (owner_cpf and u_cpf and owner_cpf == u_cpf)
        or (not owner_cpf and owner_email and u_email and owner_email == u_email)
    )


def _can_access_submission(row: Dict[str, Any], user: Dict[str, Any]) -> bool:
    """
    Permite acesso ao dono da submissão ou a papéis elevados.
    """
    if _owns_submission(row, user):
        return True
    if _has_any_role(user, *ELEVATED_ROLES):
        return True
    #if _has_any_role(user, *REQUIRED_ROLES):
    #    return True
    return False


def _read_html(name: str, fallback_title: str) -> str:
    """
    Lê HTML da pasta da automação. Se ainda não existir, retorna placeholder amigável.
    """
    path = TPL_DIR / name
    if path.exists():
        return path.read_text(encoding="utf-8")

    return f"""<!doctype html>
<html lang="pt-BR">
<meta charset="utf-8"/>
<title>{fallback_title}</title>
<div style="font-family:system-ui;padding:24px;max-width:900px;margin:0 auto">
  <div style="border:1px solid #e2e8f0;border-radius:14px;padding:16px;background:#fff">
    <h1 style="margin:0 0 8px;font-size:20px">{fallback_title}</h1>
    <p style="margin:0;color:#334155;line-height:1.5">
      A interface HTML desta automação ainda não foi criada em
      <code>apps/bff/app/automations/templates/etp/{name}</code>.
    </p>
  </div>
</div>
</html>"""


def _pydantic_errors(exc: Exception) -> List[Dict[str, Any]]:
    """
    Extrai erros amigáveis do Pydantic v2.
    """
    if not hasattr(exc, "errors"):
        return [{"field": None, "message": str(exc)}]

    out: List[Dict[str, Any]] = []
    for err in exc.errors():  # type: ignore[attr-defined]
        loc = ".".join(str(x) for x in err.get("loc", []))
        out.append({"field": loc or None, "message": err.get("msg")})
    return out


class EtpQuantitativoItem(BaseModel):
    """
    Linha da tabela de estimativas das quantidades.
    """
    model_config = ConfigDict(populate_by_name=True, extra="ignore")

    item: str = Field(..., min_length=1, max_length=20)
    especificacao: str = Field(..., min_length=1, max_length=MAX_LONG)
    unidade: str = Field(..., min_length=1, max_length=MAX_SHORT)
    quantidade: str = Field(..., min_length=1, max_length=100)

    @field_validator("item", "especificacao", "unidade", "quantidade", mode="before")
    @classmethod
    def _strip_text(cls, v: Any) -> str:
        return _normalize_text(v)


class ETPIn(BaseModel):
    """
    Payload principal do ETP.
    """
    model_config = ConfigDict(populate_by_name=True, extra="ignore")

    protocolo: Optional[str] = Field(default=None, max_length=100)
    objeto: str = Field(..., min_length=1, max_length=MAX_MEDIUM)
    unidade_demandante: str = Field(..., alias="unidadeDemandante", min_length=1, max_length=MAX_SHORT)
    responsavel: str = Field(..., min_length=1, max_length=MAX_SHORT)
    email: Optional[str] = Field(default=None, max_length=200)
    telefone: Optional[str] = Field(default=None, max_length=100)

    descricao_necessidade: str = Field(..., alias="descricaoNecessidade", min_length=1, max_length=MAX_LONG)
    previsao_pca: Optional[str] = Field(default=None, alias="previsaoPca", max_length=MAX_LONG)
    requisitos_contratacao: str = Field(..., alias="requisitosContratacao", min_length=1, max_length=MAX_LONG)

    memoria_calculo: Optional[str] = Field(default=None, alias="memoriaCalculo", max_length=MAX_LONG)
    itens_quantitativos: List[EtpQuantitativoItem] = Field(default_factory=list, alias="itensQuantitativos")

    solucoes_existentes: List[str] = Field(default_factory=list, alias="solucoesExistentes")
    analise_solucoes: List[str] = Field(default_factory=list, alias="analiseSolucoes")
    conclusao_solucao: str = Field(..., alias="conclusaoSolucao", min_length=1, max_length=MAX_LONG)

    estimativa_valor: Optional[str] = Field(default=None, alias="estimativaValor", max_length=MAX_LONG)
    descricao_solucao_como_um_todo: Optional[str] = Field(
        default=None, alias="descricaoSolucaoComoUmTodo", max_length=MAX_LONG
    )
    justificativa_parcelamento: Optional[str] = Field(
        default=None, alias="justificativaParcelamento", max_length=MAX_LONG
    )
    resultados_pretendidos: Optional[str] = Field(default=None, alias="resultadosPretendidos", max_length=MAX_LONG)
    providencias_previas: Optional[str] = Field(default=None, alias="providenciasPrevias", max_length=MAX_LONG)
    contratacoes_correlatas: Optional[str] = Field(
        default=None, alias="contratacoesCorrelatas", max_length=MAX_LONG
    )
    impactos_ambientais: Optional[str] = Field(default=None, alias="impactosAmbientais", max_length=MAX_LONG)
    posicionamento_conclusivo: str = Field(
        ..., alias="posicionamentoConclusivo", min_length=1, max_length=MAX_LONG
    )

    data_assinatura: Optional[str] = Field(default=None, alias="dataAssinatura", max_length=50)
    responsaveis_elaboracao: List[str] = Field(default_factory=list, alias="responsaveisElaboracao")
    diretor_area_demandante: Optional[str] = Field(default=None, alias="diretorAreaDemandante", max_length=MAX_SHORT)
    aprovacao_autoridade_maxima: Optional[str] = Field(
        default=None, alias="aprovacaoAutoridadeMaxima", max_length=MAX_LONG
    )

    fase_riscos: Optional[str] = Field(default=None, alias="faseRiscos", max_length=MAX_LONG)
    mapa_riscos: Optional[str] = Field(default=None, alias="mapaRiscos", max_length=MAX_LONG)

    @field_validator(
        "protocolo",
        "objeto",
        "unidade_demandante",
        "responsavel",
        "email",
        "telefone",
        "descricao_necessidade",
        "previsao_pca",
        "requisitos_contratacao",
        "memoria_calculo",
        "conclusao_solucao",
        "estimativa_valor",
        "descricao_solucao_como_um_todo",
        "justificativa_parcelamento",
        "resultados_pretendidos",
        "providencias_previas",
        "contratacoes_correlatas",
        "impactos_ambientais",
        "posicionamento_conclusivo",
        "data_assinatura",
        "diretor_area_demandante",
        "aprovacao_autoridade_maxima",
        "fase_riscos",
        "mapa_riscos",
        mode="before",
    )
    @classmethod
    def _strip_optional_text(cls, v: Any) -> Optional[str]:
        txt = _normalize_text(v)
        return txt or None

    @field_validator("solucoes_existentes", "analise_solucoes", "responsaveis_elaboracao", mode="before")
    @classmethod
    def _coerce_string_lists(cls, v: Any) -> List[str]:
        return _normalize_string_list(v)

    @field_validator("email")
    @classmethod
    def _validate_email(cls, v: Optional[str]) -> Optional[str]:
        if v and not EMAIL_RE.match(v):
            raise ValueError("E-mail inválido.")
        return v

    @model_validator(mode="after")
    def _after_model(self) -> "ETPIn":
        if not self.itens_quantitativos:
            raise ValueError("Informe ao menos um item na tabela de estimativas das quantidades.")
        if not self.solucoes_existentes:
            raise ValueError("Informe ao menos uma solução existente no mercado.")
        if not self.analise_solucoes:
            raise ValueError("Informe ao menos uma análise das soluções existentes.")
        if not self.responsaveis_elaboracao:
            self.responsaveis_elaboracao = [self.responsavel]
        return self


def _normalize_submit_body(body: Dict[str, Any]) -> Dict[str, Any]:
    """
    Normaliza o payload bruto aceitando aliases simples da UI.
    """
    solucoes_existentes = _coalesce_string_list(
        body,
        "solucoesExistentes",
        "solucoesMercado",
        "solucoes",
        "descricaoSolucoes",
    )
    if not solucoes_existentes:
        for key in sorted(body.keys()):
            if LEGACY_SOLUCAO_FIELD_RE.match(key):
                txt = _normalize_text(body.get(key))
                if txt:
                    solucoes_existentes.append(txt)

    analise_solucoes = _coalesce_string_list(
        body,
        "analiseSolucoes",
        "analisesSolucoes",
        "analises",
    )
    if not analise_solucoes:
        for key in sorted(body.keys()):
            if LEGACY_ANALISE_FIELD_RE.match(key):
                txt = _normalize_text(body.get(key))
                if txt:
                    analise_solucoes.append(txt)

    responsaveis_elaboracao = _coalesce_string_list(
        body,
        "responsaveisElaboracao",
        "responsaveis",
        "elaboradores",
    )

    raw: Dict[str, Any] = {
        "protocolo": _normalize_text(body.get("protocolo")) or None,
        "objeto": _normalize_text(body.get("objeto")),
        "unidadeDemandante": _normalize_text(body.get("unidadeDemandante") or body.get("unidade_demandante")),
        "responsavel": _normalize_text(body.get("responsavel")),
        "email": _normalize_text(body.get("email")) or None,
        "telefone": _normalize_text(body.get("telefone")) or None,
        "descricaoNecessidade": _normalize_text(
            body.get("descricaoNecessidade") or body.get("necessidade") or body.get("motivacao")
        ),
        "previsaoPca": _normalize_text(body.get("previsaoPca") or body.get("alinhamentoPca")) or None,
        "requisitosContratacao": _normalize_text(
            body.get("requisitosContratacao") or body.get("requisitos")
        ),
        "memoriaCalculo": _normalize_text(
            body.get("memoriaCalculo")
            or body.get("estimativasQuantidadeMemoriaCalculo")
            or body.get("memoria_de_calculo")
        )
        or None,
        "itensQuantitativos": _normalize_quantitativos(
            body.get("itensQuantitativos") or body.get("quantitativos") or body.get("itens")
        ),
        "solucoesExistentes": solucoes_existentes,
        "analiseSolucoes": analise_solucoes,
        "conclusaoSolucao": _normalize_text(body.get("conclusaoSolucao") or body.get("conclusao")),
        "estimativaValor": _normalize_text(body.get("estimativaValor") or body.get("valorEstimado")) or None,
        "descricaoSolucaoComoUmTodo": _normalize_text(
            body.get("descricaoSolucaoComoUmTodo") or body.get("descricaoSolucao") or body.get("objetoTecnico")
        )
        or None,
        "justificativaParcelamento": _normalize_text(
            body.get("justificativaParcelamento") or body.get("parcelamento")
        )
        or None,
        "resultadosPretendidos": _normalize_text(body.get("resultadosPretendidos")) or None,
        "providenciasPrevias": _normalize_text(body.get("providenciasPrevias")) or None,
        "contratacoesCorrelatas": _normalize_text(body.get("contratacoesCorrelatas")) or None,
        "impactosAmbientais": _normalize_text(body.get("impactosAmbientais")) or None,
        "posicionamentoConclusivo": _normalize_text(body.get("posicionamentoConclusivo")),
        "dataAssinatura": _normalize_text(body.get("dataAssinatura")) or None,
        "responsaveisElaboracao": responsaveis_elaboracao,
        "diretorAreaDemandante": _normalize_text(body.get("diretorAreaDemandante")) or None,
        "aprovacaoAutoridadeMaxima": _normalize_text(body.get("aprovacaoAutoridadeMaxima")) or None,
        "faseRiscos": _normalize_text(body.get("faseRiscos")) or None,
        "mapaRiscos": _normalize_text(body.get("mapaRiscos")) or None,
    }

    return raw


def _today_br() -> str:
    """
    Data atual no formato DD/MM/AAAA.
    """
    return datetime.utcnow().strftime("%d/%m/%Y")


def _build_context(body: ETPIn) -> Dict[str, Any]:
    """
    Produz um contexto rico para renderização DOCX.
    """
    items = [item.model_dump(by_alias=True, exclude_none=True) for item in body.itens_quantitativos]
    solucoes = list(body.solucoes_existentes or [])
    analises = list(body.analise_solucoes or [])
    responsaveis = list(body.responsaveis_elaboracao or [])

    protocolo = body.protocolo or ""
    data_assinatura = body.data_assinatura or _today_br()

    ctx: Dict[str, Any] = {
        "kind": KIND,
        "version": ETP_VERSION,
        "titulo": "ESTUDO TÉCNICO PRELIMINAR – ETP",
        "protocolo": protocolo,
        "PROTOCOLO": protocolo,
        "objeto": body.objeto,
        "OBJETO": body.objeto,
        "unidade_demandante": body.unidade_demandante,
        "unidadeDemandante": body.unidade_demandante,
        "UNIDADE_DEMANDANTE": body.unidade_demandante,
        "responsavel": body.responsavel,
        "RESPONSAVEL": body.responsavel,
        "email": body.email or "",
        "EMAIL": body.email or "",
        "telefone": body.telefone or "",
        "TELEFONE": body.telefone or "",
        "descricao_necessidade": body.descricao_necessidade,
        "descricaoNecessidade": body.descricao_necessidade,
        "previsao_pca": body.previsao_pca or "",
        "previsaoPca": body.previsao_pca or "",
        "requisitos_contratacao": body.requisitos_contratacao,
        "requisitosContratacao": body.requisitos_contratacao,
        "memoria_calculo": body.memoria_calculo or "",
        "memoriaCalculo": body.memoria_calculo or "",
        "itens_quantitativos": items,
        "itensQuantitativos": items,
        "solucoes_existentes": solucoes,
        "solucoesExistentes": solucoes,
        "solucoes_existentes_texto": "\n".join(solucoes),
        "analise_solucoes": analises,
        "analiseSolucoes": analises,
        "analise_solucoes_texto": "\n".join(analises),
        "conclusao_solucao": body.conclusao_solucao,
        "conclusaoSolucao": body.conclusao_solucao,
        "estimativa_valor": body.estimativa_valor or "",
        "estimativaValor": body.estimativa_valor or "",
        "descricao_solucao_como_um_todo": body.descricao_solucao_como_um_todo or "",
        "descricaoSolucaoComoUmTodo": body.descricao_solucao_como_um_todo or "",
        "justificativa_parcelamento": body.justificativa_parcelamento or "",
        "justificativaParcelamento": body.justificativa_parcelamento or "",
        "resultados_pretendidos": body.resultados_pretendidos or "",
        "resultadosPretendidos": body.resultados_pretendidos or "",
        "providencias_previas": body.providencias_previas or "",
        "providenciasPrevias": body.providencias_previas or "",
        "contratacoes_correlatas": body.contratacoes_correlatas or "",
        "contratacoesCorrelatas": body.contratacoes_correlatas or "",
        "impactos_ambientais": body.impactos_ambientais or "",
        "impactosAmbientais": body.impactos_ambientais or "",
        "posicionamento_conclusivo": body.posicionamento_conclusivo,
        "posicionamentoConclusivo": body.posicionamento_conclusivo,
        "data_assinatura": data_assinatura,
        "dataAssinatura": data_assinatura,
        "responsaveis_elaboracao": responsaveis,
        "responsaveisElaboracao": responsaveis,
        "responsaveis_elaboracao_texto": "\n".join(responsaveis),
        "diretor_area_demandante": body.diretor_area_demandante or "",
        "diretorAreaDemandante": body.diretor_area_demandante or "",
        # Campo mantido em comentário para possível reativação futura,
        # mas oculto da geração atual do documento.
        # "aprovacao_autoridade_maxima": body.aprovacao_autoridade_maxima or "",
        # "aprovacaoAutoridadeMaxima": body.aprovacao_autoridade_maxima or "",
        "fase_riscos": body.fase_riscos or "",
        "faseRiscos": body.fase_riscos or "",
        "mapa_riscos": body.mapa_riscos or "",
        "mapaRiscos": body.mapa_riscos or "",
        "gerado_em": _today_br(),
    }
    return ctx


def _apply_run_font(paragraph, *, bold: bool = False, size: int = 11) -> None:
    """
    Ajusta formatação básica do último run do parágrafo.
    """
    if not paragraph.runs:
        return
    run = paragraph.runs[-1]
    run.bold = bold
    if Pt is not None:
        run.font.size = Pt(size)


def _add_paragraph(doc, text: str = "", *, bold: bool = False, size: int = 11, align: Optional[int] = None):
    """
    Cria um parágrafo com configuração simples.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6) if Pt is not None else None
    if align is not None and WD_ALIGN_PARAGRAPH is not None:
        p.alignment = align
    if text:
        p.add_run(text)
        _apply_run_font(p, bold=bold, size=size)
    return p


def _add_multiline_block(doc, text: str) -> None:
    """
    Adiciona um bloco de texto respeitando quebras de linha.
    """
    content = _normalize_text(text)
    if not content:
        _add_paragraph(doc, "Não informado.")
        return
    for line in content.replace("\r", "\n").split("\n"):
        line = line.strip()
        if line:
            _add_paragraph(doc, line)
    if "\n" not in content:
        return


def _add_section(doc, title: str, text: str) -> None:
    """
    Adiciona uma seção textual do ETP.
    """
    _add_paragraph(doc, title, bold=True, size=12)
    _add_multiline_block(doc, text)
    _add_paragraph(doc)


def _add_string_list(doc, items: List[str], *, prefix: str) -> None:
    """
    Adiciona uma sequência enumerada de blocos textuais preservando
    quebras de linha dentro de cada item.
    """
    if not items:
        _add_paragraph(doc, "Não informado.")
        return

    for idx, item in enumerate(items, start=1):
        content = _normalize_text(item)
        if not content:
            continue

        lines = [line.strip() for line in content.replace("\r", "\n").split("\n") if line.strip()]
        if not lines:
            continue

        _add_paragraph(doc, f"{prefix} {idx}: {lines[0]}")
        for line in lines[1:]:
            _add_paragraph(doc, line)


def _clear_document_body(doc) -> None:
    """
    Remove o conteúdo do corpo preservando `sectPr` e, por consequência,
    cabeçalho, rodapé, margens e configuração de página.
    """
    if qn is None:
        return
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _render_etp_docx_blank_template(template_path: str, context: Dict[str, Any], out_path: str) -> None:
    """
    Renderiza o ETP em um DOCX em branco com timbre, preservando cabeçalho e rodapé.
    """
    if Document is None:
        raise RuntimeError("python-docx não está disponível no ambiente do BFF.")

    doc = Document(template_path)
    _clear_document_body(doc)

    _add_paragraph(
        doc,
        context.get("titulo") or "ESTUDO TÉCNICO PRELIMINAR – ETP",
        bold=True,
        size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER if WD_ALIGN_PARAGRAPH is not None else None,
    )
    _add_paragraph(doc)

    _add_paragraph(doc, f"PROTOCOLO: {context.get('protocolo') or '—'}", bold=True)
    _add_paragraph(doc, f"OBJETO: {context.get('objeto') or '—'}")
    _add_paragraph(doc, f"UNIDADE DEMANDANTE: {context.get('unidade_demandante') or '—'}")
    _add_paragraph(doc, f"RESPONSÁVEL: {context.get('responsavel') or '—'}")
    _add_paragraph(doc, f"E-MAIL: {context.get('email') or '—'}")
    _add_paragraph(doc, f"TELEFONE: {context.get('telefone') or '—'}")
    _add_paragraph(doc)

    _add_section(doc, "1. Descrição da necessidade da contratação", context.get("descricao_necessidade") or "")
    _add_section(
        doc,
        "2. Previsão da contratação no Plano de Contratações Anual – PCA, sempre que elaborado",
        context.get("previsao_pca") or "",
    )
    _add_section(doc, "3. Requisitos da contratação", context.get("requisitos_contratacao") or "")

    _add_paragraph(doc, "4. Estimativas das quantidades para a contratação | Memória de cálculo", bold=True, size=12)
    _add_multiline_block(doc, context.get("memoria_calculo") or "")
    _add_paragraph(doc)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Especificação"
    hdr[2].text = "Unidade"
    hdr[3].text = "Quantidade"

    for row in context.get("itens_quantitativos") or []:
        cells = table.add_row().cells
        cells[0].text = _normalize_text(row.get("item"))
        cells[1].text = _normalize_text(row.get("especificacao"))
        cells[2].text = _normalize_text(row.get("unidade"))
        cells[3].text = _normalize_text(row.get("quantidade"))

    _add_paragraph(doc)

    _add_paragraph(
        doc,
        "5. Levantamento de mercado: análise das soluções existentes no mercado e justificativa técnica e econômica da escolha do tipo de solução a contratar",
        bold=True,
        size=12,
    )
    _add_paragraph(doc, "5.1 Soluções existentes no mercado", bold=True)
    _add_string_list(doc, context.get("solucoes_existentes") or [], prefix="Descrição da Solução")
    _add_paragraph(doc)
    _add_paragraph(doc, "5.2 Análise das soluções existentes", bold=True)
    _add_string_list(doc, context.get("analise_solucoes") or [], prefix="Análise da Solução")
    _add_paragraph(doc)
    _add_paragraph(doc, "5.3 Conclusão quanto à solução a ser adotada e os motivos da escolha", bold=True)
    _add_multiline_block(doc, context.get("conclusao_solucao") or "")
    _add_paragraph(doc)

    _add_section(doc, "6. Estimativa do valor da contratação", context.get("estimativa_valor") or "")
    _add_section(doc, "7. Descrição da solução como um todo – Objeto Técnico", context.get("descricao_solucao_como_um_todo") or "")
    _add_section(doc, "8. Justificativas para o parcelamento ou não da contratação", context.get("justificativa_parcelamento") or "")
    _add_section(
        doc,
        "9. Demonstrativo dos resultados pretendidos em termos de economicidade e de melhor aproveitamento dos recursos humanos, materiais e financeiros disponíveis",
        context.get("resultados_pretendidos") or "",
    )
    _add_section(
        doc,
        "10. Providências a serem adotadas pela Administração previamente à celebração do contrato",
        context.get("providencias_previas") or "",
    )
    _add_section(doc, "11. Contratações correlatas e/ou interdependentes", context.get("contratacoes_correlatas") or "")
    _add_section(doc, "12. Descrição de possíveis impactos ambientais e respectivas medidas mitigadoras", context.get("impactos_ambientais") or "")
    _add_section(
        doc,
        "13. Posicionamento conclusivo sobre a adequação da contratação para o atendimento da necessidade a que se destina",
        context.get("posicionamento_conclusivo") or "",
    )

    _add_paragraph(doc, "14. Data e assinatura do(s) responsável(is) pela elaboração do ETP e Diretor(a) da área demandante", bold=True, size=12)
    _add_paragraph(doc, f"Data: {context.get('data_assinatura') or _today_br()}")
    _add_paragraph(doc, "Responsável(is) pela elaboração:", bold=True)
    responsaveis = context.get("responsaveis_elaboracao") or []
    if responsaveis:
        for nome in responsaveis:
            _add_paragraph(doc, f"- {nome}")
    else:
        _add_paragraph(doc, "- Não informado.")
    _add_paragraph(doc, f"Diretor(a) da área demandante: {context.get('diretor_area_demandante') or 'Não informado.'}")
    _add_paragraph(doc)

    # Seção mantida em comentário para possível reativação futura,
    # mas não deve aparecer no documento nesta fase.
    # _add_section(doc, "15. Aprovação da autoridade máxima da Agepar", context.get("aprovacao_autoridade_maxima") or "")
    _add_section(doc, "16. Fase de identificação e análise dos riscos", context.get("fase_riscos") or "")
    _add_section(doc, "16.1 Mapa de Riscos", context.get("mapa_riscos") or "")

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)


def _render_etp_docx(template_path: str, context: Dict[str, Any], out_path: str) -> None:
    """
    Usa docxtpl quando houver placeholders; caso contrário, monta o corpo manualmente.
    """
    vars_ = get_docx_placeholders(template_path)
    if vars_:
        render_docx_template(template_path, context, out_path)
        return
    _render_etp_docx_blank_template(template_path, context, out_path)


def _process_submission(sid: str, body: ETPIn, actor: Dict[str, Any]) -> None:
    """
    Pipeline assíncrono:
    1) marca `running`;
    2) gera DOCX (e PDF quando possível);
    3) persiste o resultado final em `submissions.result`;
    4) audita sucesso/falha.
    """
    try:
        update_submission(sid, status="running")
        add_audit(KIND, "running", actor, {"sid": sid})
    except Exception as e:
        logger.exception("[ETP] update to running failed")
        try:
            update_submission(sid, status="error", error=f"storage: {e}")
        except Exception:
            pass
        try:
            add_audit(KIND, "failed", actor, {"sid": sid, "error": f"storage: {e}"})
        except Exception:
            pass
        return

    try:
        if not TEMPLATE_PATH.exists():
            raise RuntimeError(f"Template não encontrado em {TEMPLATE_PATH}")

        context = _build_context(body)

        out_dir = OUTPUT_DIR / sid
        out_dir.mkdir(parents=True, exist_ok=True)

        base_hint = _safe_comp(body.protocolo or body.objeto or sid)
        base_name = f"etp_{base_hint}_{datetime.utcnow().date().isoformat()}"
        docx_filename = f"{base_name}.docx"
        pdf_filename = f"{base_name}.pdf"

        docx_path = str(out_dir / docx_filename)
        pdf_path = str(out_dir / pdf_filename)

        _render_etp_docx(str(TEMPLATE_PATH), context, docx_path)

        final_path = docx_path
        final_name = docx_filename
        pdf_generated = False

        try:
            pdf_generated = convert_docx_to_pdf(docx_path, pdf_path)
        except Exception:
            logger.exception("[ETP] convert_docx_to_pdf falhou; mantendo DOCX")

        if pdf_generated and os.path.exists(pdf_path):
            final_path = pdf_path
            final_name = pdf_filename

        result = {
            "sid": sid,
            "file_path": final_path,
            "filename": final_name,
            "docx_path": docx_path,
            "docx_filename": docx_filename,
            "pdf_path": pdf_path if os.path.exists(pdf_path) else None,
            "pdf_filename": pdf_filename if os.path.exists(pdf_path) else None,
            "template_path": str(TEMPLATE_PATH),
            "generated_at": datetime.utcnow().isoformat() + "Z",
        }

        update_submission(sid, status="done", result=result, error=None)
        add_audit(KIND, "completed", actor, {"sid": sid, "filename": final_name})
        logger.info("[ETP] Submissão %s finalizada | entregue=%s", sid, final_name)

    except Exception as e:
        logger.exception("[ETP] processing error")
        try:
            update_submission(sid, status="error", error=str(e))
        except Exception:
            pass
        try:
            add_audit(KIND, "failed", actor, {"sid": sid, "error": str(e)})
        except Exception:
            pass


@router.get("/schema")
async def schema_etp(
    user: Dict[str, Any] = Depends(require_roles_any(*REQUIRED_ROLES)),
):
    """
    Metadados do formulário ETP.
    """
    return {
        "kind": KIND,
        "version": ETP_VERSION,
        "requiredRoles": list(REQUIRED_ROLES),
        "templatePath": str(TEMPLATE_PATH),
        "templateExists": TEMPLATE_PATH.exists(),
        "schema": ETPIn.model_json_schema(by_alias=True),
    }


@router.get("/submissions")
async def list_my_submissions(
    user: Dict[str, Any] = Depends(require_roles_any(*REQUIRED_ROLES)),
    limit: int = 50,
    offset: int = 0,
):
    """
    Lista submissões do usuário autenticado.
    """
    try:
        cpf = (user.get("cpf") or "").strip() or None
        email = (user.get("email") or "").strip() or None
        if not cpf and not email:
            return err_json(400, code="missing_identity", message="Usuário sem identificador (cpf/email).")

        rows = list_submissions(
            kind=KIND,
            actor_cpf=cpf,
            actor_email=None if cpf else email,
            limit=limit,
            offset=offset,
        )
        return {"items": rows, "limit": limit, "offset": offset}
    except Exception as e:
        logger.exception("[ETP] list_submissions failed")
        return err_json(500, code="storage_error", message="Falha ao consultar submissões.", details=str(e))


@router.get("/submissions/{sid}")
async def get_my_submission(
    sid: str,
    user: Dict[str, Any] = Depends(require_roles_any(*REQUIRED_ROLES)),
):
    """
    Retorna uma submissão do usuário, com checagem de ownership.
    """
    try:
        row = get_submission(sid)
    except Exception as e:
        logger.exception("[ETP] get_submission failed")
        return err_json(500, code="storage_error", message="Falha ao consultar submissão.", details=str(e))

    if not row:
        return err_json(404, code="not_found", message="Submissão não encontrada.", details={"sid": sid})
    if not _owns_submission(row, user):
        return err_json(403, code="forbidden", message="Você não tem acesso a esta submissão.")
    return row


@router.post("/submit")
async def submit_etp(
    request: Request,
    body: Dict[str, Any],
    background: BackgroundTasks,
    user: Dict[str, Any] = Depends(require_roles_any(*REQUIRED_ROLES)),
):
    """
    Recebe o formulário do ETP, normaliza, valida e agenda a geração do documento.
    """
    raw = _normalize_submit_body(body)

    if not raw.get("objeto"):
        return err_json(422, code="validation_error", message="Objeto é obrigatório.")
    if not raw.get("unidadeDemandante"):
        return err_json(422, code="validation_error", message="Unidade demandante é obrigatória.")
    if not raw.get("responsavel"):
        return err_json(422, code="validation_error", message="Responsável é obrigatório.")
    if not raw.get("descricaoNecessidade"):
        return err_json(422, code="validation_error", message="Descrição da necessidade é obrigatória.")

    try:
        payload = ETPIn.model_validate(raw)
    except ValidationError as ve:
        friendly = _pydantic_errors(ve)
        logger.info("[ETP] validation_error: %s", friendly)
        return err_json(
            422,
            code="validation_error",
            message="Erro de validação nos campos.",
            details={"errors": friendly},
        )
    except Exception as e:
        logger.exception("[ETP] validation failure")
        return err_json(422, code="validation_error", message="Erro de validação.", details=str(e))

    sid = str(uuid4())
    sub = {
        "id": sid,
        "kind": KIND,
        "version": ETP_VERSION,
        "actor_cpf": user.get("cpf"),
        "actor_nome": user.get("nome"),
        "actor_email": user.get("email"),
        "payload": payload.model_dump(by_alias=True, exclude_none=True),
        "status": "queued",
        "result": None,
        "error": None,
    }

    try:
        insert_submission(sub)
        add_audit(KIND, "submitted", user, {"sid": sid, "protocolo": payload.protocolo})
    except Exception as e:
        logger.exception("[ETP] insert_submission failed")
        return err_json(500, code="storage_error", message="Falha ao salvar a submissão.", details=str(e))

    # Notifica usuários com cargo/role CA (best-effort; não pode bloquear o ETP).
    try:
        protocolo_val = (payload.protocolo or "").strip()
        objeto_val = (payload.objeto or "").strip()

        msg_parts = ["Um novo ETP foi enviado."]
        if protocolo_val:
            msg_parts.append(f"Protocolo: {protocolo_val}.")
        if objeto_val:
            msg_parts.append(f"Objeto: {objeto_val}.")
        msg_parts.append(f"Autor: {user.get('nome') or '—'}")

        notif_id, delivered = send_notification(
            actor=user,
            title="Novo ETP enviado",
            message=" ".join(msg_parts),
            role_names=["coordendor_ca"],
            action_url="/controle",
            meta={"kind": "etp", "submissionId": sid, "protocolo": protocolo_val, "objeto": objeto_val},
            ip=request.client.host if request.client else None,
            ua=request.headers.get("user-agent"),
        )
        logger.info("[ETP] Notificação enviada para COORDENDOR_CA | notif=%s | delivered=%d", notif_id, delivered)
    except Exception:
        logger.exception("[ETP] Falha ao notificar COORDENDOR_CA (não bloqueante)")

    logger.info(
        "[ETP] Submissão %s criada por %s (%s) | protocolo=%s",
        sid,
        user.get("nome"),
        user.get("cpf"),
        payload.protocolo or "—",
    )

    background.add_task(_process_submission, sid, payload, user)
    return {"submissionId": sid, "status": "queued"}


@router.post("/submissions/{sid}/download")
async def download_result(
    sid: str,
    request: Request,
    user: Dict[str, Any] = Depends(require_roles_any("compras", "admin")),
):
    """
    Download primário: retorna PDF quando disponível; caso contrário, DOCX.
    """
    try:
        row = get_submission(sid)
    except Exception as e:
        logger.exception("[ETP] get_submission(download) failed")
        return err_json(500, code="storage_error", message="Falha ao consultar submissão.", details=str(e))

    if not row:
        return err_json(404, code="not_found", message="Submissão não encontrada.", details={"sid": sid})
    if not _can_access_submission(row, user):
        return err_json(403, code="forbidden", message="Você não tem acesso a esta submissão.")
    if row.get("status") != "done":
        return err_json(
            409,
            code="not_ready",
            message="Resultado ainda não está pronto.",
            details={"status": row.get("status")},
        )

    try:
        result = _to_obj(row.get("result"), {})
        file_path = result.get("file_path")
        filename = result.get("filename") or f"etp_{sid}.docx"

        if not file_path or not os.path.exists(file_path):
            return err_json(410, code="file_not_found", message="Arquivo não está mais disponível.", details={"sid": sid})

        with open(file_path, "rb") as f:
            data = f.read()

        try:
            ext = (os.path.splitext(filename)[1] or "").lstrip(".").lower() or "auto"
            add_audit(
                KIND,
                "download",
                user,
                {
                    "sid": sid,
                    "filename": filename,
                    "bytes": len(data),
                    "fmt": ext,
                    "ip": (getattr(request.client, "host", None) if request and request.client else None),
                    "ua": (request.headers.get("user-agent") if request else None),
                },
            )
        except Exception:
            logger.exception("[ETP] audit(download) failed (non-blocking)")

        media_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"
        return StreamingResponse(
            BytesIO(data),
            media_type=media_type,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        logger.exception("[ETP] download error")
        return err_json(500, code="download_error", message="Falha ao preparar o download.", details=str(e))


@router.post("/submissions/{sid}/download/{fmt}")
async def download_result_fmt(
    sid: str,
    fmt: str,
    request: Request,
    user: Dict[str, Any] = Depends(require_roles_any("compras", "admin")),
):
    """
    Download explícito por formato: `pdf` ou `docx`.
    """
    if fmt not in ("pdf", "docx"):
        return err_json(400, code="bad_request", message="Formato inválido. Use 'pdf' ou 'docx'.")

    try:
        row = get_submission(sid)
    except Exception as e:
        logger.exception("[ETP] get_submission(download fmt) failed")
        return err_json(500, code="storage_error", message="Falha ao consultar submissão.", details=str(e))

    if not row:
        return err_json(404, code="not_found", message="Submissão não encontrada.", details={"sid": sid})
    if not _can_access_submission(row, user):
        return err_json(403, code="forbidden", message="Você não tem acesso a esta submissão.")
    if row.get("status") != "done":
        return err_json(
            409,
            code="not_ready",
            message="Resultado ainda não está pronto.",
            details={"status": row.get("status")},
        )

    try:
        result = _to_obj(row.get("result"), {})
        file_path = result.get(f"{fmt}_path")
        filename = result.get(f"{fmt}_filename")

        if not file_path or not filename:
            return err_json(404, code="format_not_available", message=f"Arquivo {fmt.upper()} não disponível.")
        if not os.path.exists(file_path):
            return err_json(410, code="file_not_found", message="Arquivo não está mais disponível.", details={"sid": sid})

        with open(file_path, "rb") as f:
            data = f.read()

        try:
            add_audit(
                KIND,
                "download",
                user,
                {
                    "sid": sid,
                    "filename": filename,
                    "bytes": len(data),
                    "fmt": fmt,
                    "ip": (getattr(request.client, "host", None) if request and request.client else None),
                    "ua": (request.headers.get("user-agent") if request else None),
                },
            )
        except Exception:
            logger.exception("[ETP] audit(download fmt) failed (non-blocking)")

        media_type = mimetypes.guess_type(filename)[0] or "application/octet-stream"
        return StreamingResponse(
            BytesIO(data),
            media_type=media_type,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    except Exception as e:
        logger.exception("[ETP] download fmt error")
        return err_json(500, code="download_error", message="Falha ao preparar o download.", details=str(e))


@router.get("/audits")
async def list_audits_admin(
    user: Dict[str, Any] = Depends(require_roles_any("admin")),
    limit: int = 50,
    offset: int = 0,
):
    """
    Lista auditorias da automação ETP.
    """
    try:
        rows = list_audits(kind=KIND, limit=limit, offset=offset)
        return {"items": rows, "limit": limit, "offset": offset}
    except Exception as e:
        logger.exception("[ETP] list_audits failed")
        return err_json(500, code="storage_error", message="Falha ao consultar auditoria.", details=str(e))


@router.get("/people/search")
async def search_people(
    q: str = "",
    limit: int = 10,
    user: Dict[str, Any] = Depends(require_roles_any(*REQUIRED_ROLES)),
):
    """
    Busca simples de pessoas cadastradas para autocomplete.
    - Retorna matches por nome, e-mail ou CPF.
    - Não bloqueia texto livre na UI; serve apenas como sugestão.
    """
    term = (q or "").strip()
    lim = max(1, min(int(limit or 10), 20))

    if len(term) < 2:
        return {"items": []}

    if not DATABASE_URL:
        return {"items": []}

    try:
        with psycopg.connect(DATABASE_URL, autocommit=True) as conn:
            with conn.cursor() as cur:
                like = f"%{term}%"
                cur.execute(
                    """
                    SELECT
                        id::text,
                        COALESCE(name, '') AS nome,
                        COALESCE(email_institucional::text, email::text, '') AS email,
                        COALESCE(cpf, '') AS cpf
                    FROM users
                    WHERE
                        COALESCE(name, '') ILIKE %s
                        OR COALESCE(email::text, '') ILIKE %s
                        OR COALESCE(email_institucional::text, '') ILIKE %s
                        OR COALESCE(cpf, '') ILIKE %s
                    ORDER BY
                        CASE
                            WHEN COALESCE(name, '') ILIKE %s THEN 0
                            ELSE 1
                        END,
                        name ASC
                    LIMIT %s
                    """,
                    (like, like, like, like, f"{term}%", lim),
                )
                rows = cur.fetchall()
        items = [{"id": r[0], "nome": r[1], "email": r[2], "cpf": r[3]} for r in rows]
        return {"items": items}
    except Exception as e:
        logger.exception("[ETP] people search failed")
        return err_json(500, code="people_search_error", message="Falha ao buscar pessoas cadastradas.", details=str(e))


@router.get("/ui")
@router.get("/ui/")
async def etp_ui(request: Request):
    """
    Página principal da UI do ETP.
    """
    checker = require_roles_any(*REQUIRED_ROLES)
    try:
        checker(request)
    except HTTPException as he:
        status = he.status_code
        msg = (
            "Faça login para acessar esta automação."
            if status == 401
            else "Você não tem permissão para acessar esta automação."
        )
        html_err = f"""<!doctype html><meta charset="utf-8"/><title>Acesso</title>
        <div style="font-family:system-ui;padding:24px">
          <h1 style="margin:0 0 8px">{status}</h1>
          <p style="color:#334155">{msg}</p>
        </div>"""
        return HTMLResponse(html_err, status_code=status)

    return HTMLResponse(_read_html("ui.html", "ETP"))


@router.get("/ui/history")
@router.get("/ui/history/")
async def etp_history_ui(request: Request):
    """
    Página de histórico do ETP.
    """
    checker = require_roles_any(*REQUIRED_ROLES)
    try:
        checker(request)
    except HTTPException as he:
        status = he.status_code
        msg = (
            "Faça login para acessar esta automação."
            if status == 401
            else "Você não tem permissão para acessar esta automação."
        )
        html_err = f"""<!doctype html><meta charset="utf-8"/><title>Acesso</title>
        <div style="font-family:system-ui;padding:24px">
          <h1 style="margin:0 0 8px">{status}</h1>
          <p style="color:#334155">{msg}</p>
        </div>"""
        return HTMLResponse(html_err, status_code=status)

    return HTMLResponse(_read_html("history.html", "Histórico do ETP"))
